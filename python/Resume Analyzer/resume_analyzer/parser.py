"""
parser.py
---------
Extracts raw text from resume files (PDF, DOCX, TXT).
"""

import os
import re
import pdfplumber
import docx


def extract_text(filepath: str) -> str:
    """Extract raw text from a resume file, regardless of format."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(filepath)
    elif ext == ".docx":
        return _extract_from_docx(filepath)
    elif ext == ".txt":
        return _extract_from_txt(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def _extract_from_pdf(filepath: str) -> str:
    text_chunks = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def _extract_from_docx(filepath: str) -> str:
    document = docx.Document(filepath)
    return "\n".join(para.text for para in document.paragraphs)


def _extract_from_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_hyperlinks(filepath: str) -> list:
    """
    Extract clickable hyperlink URLs from a resume file.

    Why this matters: many resume templates show contact info as icons
    (LinkedIn/GitHub/email icons) with NO visible text — the icon is a
    hyperlink, but plain text extraction only sees an empty glyph or
    nothing at all. Regex on extracted text alone will always say
    "not found" for those, even though the info is genuinely in the file.
    """
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return _extract_pdf_hyperlinks(filepath)
    elif ext == ".docx":
        return _extract_docx_hyperlinks(filepath)
    return []


def _extract_pdf_hyperlinks(filepath: str) -> list:
    links = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            for link in page.hyperlinks:
                uri = link.get("uri")
                if uri:
                    links.append(uri)
    return links


def _extract_docx_hyperlinks(filepath: str) -> list:
    links = []
    document = docx.Document(filepath)
    rels = document.part.rels
    for rel in rels.values():
        if "hyperlink" in rel.reltype and rel.target_ref:
            links.append(rel.target_ref)
    return links


def clean_text(raw_text: str) -> str:
    """Normalize whitespace and remove non-printable artifacts."""
    text = raw_text.replace("\x0c", "\n")  # form feed from PDFs
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_into_lines(text: str) -> list:
    """Return non-empty, stripped lines - useful for section detection."""
    return [line.strip() for line in text.split("\n") if line.strip()]
